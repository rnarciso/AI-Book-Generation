# -*- coding: utf-8 -*-
# book_creator_ai.py

import json
import os
import time
import re # Para sanitizar nomes de arquivos

# --- Biblioteca da API Gemini ---
try:
    import google.generativeai as genai
    GEMINI_API_AVAILABLE = True
except ImportError:
    GEMINI_API_AVAILABLE = False
    print("AVISO: Biblioteca 'google-generativeai' não encontrada. A funcionalidade da IA será desabilitada.")
    print("Para habilitar, instale com: pip install google-generativeai")

# --- Bibliotecas para Geração de Documentos ---
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("AVISO: Biblioteca 'python-docx' não encontrada. A geração de DOCX será desabilitada.")
    print("Para habilitar, instale com: pip install python-docx")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("AVISO: Biblioteca 'reportlab' não encontrada. A geração de PDF será desabilitada.")
    print("Para habilitar, instale com: pip install reportlab")

# --- Constantes ---
NOME_ARQUIVO_PROGRESO = "livro_em_progresso.json"
MODELOS_GEMINI_SUGERIDOS = [
    "gemini-1.5-flash-latest", # Rápido e eficiente para muitas tarefas
    "gemini-1.5-pro-latest",   # Modelo mais capaz para tarefas complexas
    "gemini-1.0-pro",          # Modelo legado, ainda útil
]

# --- Variáveis Globais de Configuração da IA ---
GEMINI_API_KEY = None
GEMINI_MODEL_NAME_GERACAO = None # Modelo para geração de capítulos
GEMINI_MODEL_NAME_FINALIZACAO = None # Modelo para o agente finalizador
GENERATIVE_MODEL_INSTANCE_GERACAO = None
GENERATIVE_MODEL_INSTANCE_FINALIZACAO = None


# --- Funções Auxiliares de Interação e UI ---

def limpar_tela():
    """Limpa o terminal para melhor visualização."""
    os.system('cls' if os.name == 'nt' else 'clear')

def exibir_cabecalho(titulo):
    """Exibe um cabeçalho formatado."""
    print("\n" + "=" * 60)
    print(f"📚 {titulo.center(56)} 📚")
    print("=" * 60 + "\n")

def exibir_status(mensagem, delay=0.5):
    """Exibe uma mensagem de status."""
    print(f"⚙️  {mensagem}")
    time.sleep(delay)

def obter_input_usuario(pergunta, tipo_esperado=str, validacao_func=None, erro_msg="Entrada inválida. Tente novamente.", opcoes_validas_lista=None):
    """Obtém input do usuário com validação opcional e lista de opções."""
    while True:
        try:
            if opcoes_validas_lista:
                print(f"   Opções disponíveis: {', '.join(opcoes_validas_lista)}")
            
            resposta = input(f"➡️  {pergunta} ").strip()

            if tipo_esperado == int:
                resposta_convertida = int(resposta)
            elif tipo_esperado == float:
                resposta_convertida = float(resposta)
            else:
                resposta_convertida = resposta

            if opcoes_validas_lista and resposta_convertida not in opcoes_validas_lista:
                print(f"❌ Opção inválida. Escolha entre: {', '.join(opcoes_validas_lista)}")
                continue

            if validacao_func:
                if validacao_func(resposta_convertida):
                    return resposta_convertida
                else:
                    print(f"❌ {erro_msg}")
            else:
                if tipo_esperado == str and not resposta_convertida: # Não permitir string vazia se não especificado
                    print("❌ A entrada não pode ser vazia.")
                else:
                    return resposta_convertida
        except ValueError:
            print(f"❌ Erro: Esperava um valor do tipo '{tipo_esperado.__name__}'.")
        except Exception as e:
            print(f"❌ Ocorreu um erro inesperado: {e}")

def obter_confirmacao(pergunta, opcoes_validas=('s', 'n', 'r', 'm', 'a')):
    """Obtém confirmação do usuário (sim/não/refazer/melhorar/adicionar)."""
    opcoes_str = "/".join(opcoes_validas)
    while True:
        resposta = input(f"➡️  {pergunta} ({opcoes_str}): ").strip().lower()
        if resposta in opcoes_validas:
            return resposta
        else:
            print(f"❌ Opção inválida. Por favor, digite uma das seguintes opções: {opcoes_str}.")

# --- Configuração e Chamada da API Gemini ---
def _criar_instancia_modelo(model_name):
    """Cria uma instância de modelo generativo com configurações de segurança."""
    try:
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]
        return genai.GenerativeModel(model_name, safety_settings=safety_settings)
    except Exception as e:
        print(f"❌ Erro ao criar instância do modelo '{model_name}': {e}")
        return None

def configurar_api_gemini():
    """Configura a API Key e os modelos Gemini a serem utilizados."""
    global GEMINI_API_KEY, GEMINI_MODEL_NAME_GERACAO, GEMINI_MODEL_NAME_FINALIZACAO
    global GENERATIVE_MODEL_INSTANCE_GERACAO, GENERATIVE_MODEL_INSTANCE_FINALIZACAO

    if not GEMINI_API_AVAILABLE:
        print("❌ API Gemini não está disponível (biblioteca não instalada). Não é possível configurar.")
        return False

    exibir_cabecalho("Configuração da API Gemini")
    print("Você precisará de uma API Key do Google AI Studio (https://aistudio.google.com/app/apikey).")
    
    if os.getenv("GEMINI_API_KEY"):
        print("ℹ️  API Key encontrada na variável de ambiente GEMINI_API_KEY.")
        GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    else:
        GEMINI_API_KEY = obter_input_usuario("Por favor, insira sua API Key do Gemini:")

    try:
        genai.configure(api_key=GEMINI_API_KEY)
        print("✅ API Key configurada com sucesso.")
    except Exception as e:
        print(f"❌ Erro ao configurar a API Key: {e}")
        print("   Por favor, verifique sua API Key e tente novamente.")
        return False

    # Seleção do modelo para GERAÇÃO
    print("\n--- Seleção de Modelo para Geração de Conteúdo (Capítulos) ---")
    print("Modelos Gemini sugeridos:")
    for i, modelo in enumerate(MODELOS_GEMINI_SUGERIDOS):
        print(f"  {i+1}. {modelo}")
    print(f"  {len(MODELOS_GEMINI_SUGERIDOS)+1}. Digitar outro nome de modelo")

    while True:
        escolha_modelo_str = obter_input_usuario(f"Escolha o modelo para GERAÇÃO (1-{len(MODELOS_GEMINI_SUGERIDOS)+1}):")
        try:
            escolha_modelo = int(escolha_modelo_str)
            if 1 <= escolha_modelo <= len(MODELOS_GEMINI_SUGERIDOS):
                GEMINI_MODEL_NAME_GERACAO = MODELOS_GEMINI_SUGERIDOS[escolha_modelo-1]
                break
            elif escolha_modelo == len(MODELOS_GEMINI_SUGERIDOS)+1:
                GEMINI_MODEL_NAME_GERACAO = obter_input_usuario("Digite o nome do modelo Gemini para GERAÇÃO (ex: 'gemini-1.5-flash-latest'):")
                break
            else:
                print(f"❌ Escolha inválida.")
        except ValueError:
            print(f"❌ Entrada inválida. Por favor, digite um número.")
            
    GENERATIVE_MODEL_INSTANCE_GERACAO = _criar_instancia_modelo(GEMINI_MODEL_NAME_GERACAO)
    if not GENERATIVE_MODEL_INSTANCE_GERACAO: return False
    print(f"✅ Modelo para GERAÇÃO '{GEMINI_MODEL_NAME_GERACAO}' selecionado.")

    # Seleção do modelo para FINALIZAÇÃO
    print("\n--- Seleção de Modelo para Finalização e Revisão Geral (Agente Finalizador) ---")
    if obter_confirmacao(f"Deseja usar um modelo DIFERENTE (potencialmente mais robusto como gemini-1.5-pro-latest) para a finalização? (s/n - para usar '{GEMINI_MODEL_NAME_GERACAO}')", ('s','n')) == 's':
        print("Modelos Gemini sugeridos (geralmente um modelo 'pro' é bom para revisão):")
        for i, modelo in enumerate(MODELOS_GEMINI_SUGERIDOS):
            print(f"  {i+1}. {modelo}")
        print(f"  {len(MODELOS_GEMINI_SUGERIDOS)+1}. Digitar outro nome de modelo")
        while True:
            escolha_modelo_str = obter_input_usuario(f"Escolha o modelo para FINALIZAÇÃO (1-{len(MODELOS_GEMINI_SUGERIDOS)+1}):")
            try:
                escolha_modelo = int(escolha_modelo_str)
                if 1 <= escolha_modelo <= len(MODELOS_GEMINI_SUGERIDOS):
                    GEMINI_MODEL_NAME_FINALIZACAO = MODELOS_GEMINI_SUGERIDOS[escolha_modelo-1]
                    break
                elif escolha_modelo == len(MODELOS_GEMINI_SUGERIDOS)+1:
                    GEMINI_MODEL_NAME_FINALIZACAO = obter_input_usuario("Digite o nome do modelo Gemini para FINALIZAÇÃO (ex: 'gemini-1.5-pro-latest'):")
                    break
                else:
                    print(f"❌ Escolha inválida.")
            except ValueError:
                print(f"❌ Entrada inválida. Por favor, digite um número.")
    else:
        GEMINI_MODEL_NAME_FINALIZACAO = GEMINI_MODEL_NAME_GERACAO
    
    GENERATIVE_MODEL_INSTANCE_FINALIZACAO = _criar_instancia_modelo(GEMINI_MODEL_NAME_FINALIZACAO)
    if not GENERATIVE_MODEL_INSTANCE_FINALIZACAO: return False
    print(f"✅ Modelo para FINALIZAÇÃO '{GEMINI_MODEL_NAME_FINALIZACAO}' selecionado.")
    
    return True


def chamar_api_gemini(prompt_para_ia, tipo_agente="geral", usar_modelo_finalizacao=False):
    """
    Chama a API Gemini real para gerar conteúdo.
    Permite escolher entre o modelo de geração e o de finalização.
    """
    global GENERATIVE_MODEL_INSTANCE_GERACAO, GENERATIVE_MODEL_INSTANCE_FINALIZACAO
    global GEMINI_MODEL_NAME_GERACAO, GEMINI_MODEL_NAME_FINALIZACAO

    modelo_a_usar = None
    nome_modelo_em_uso = ""

    if usar_modelo_finalizacao and GENERATIVE_MODEL_INSTANCE_FINALIZACAO:
        modelo_a_usar = GENERATIVE_MODEL_INSTANCE_FINALIZACAO
        nome_modelo_em_uso = GEMINI_MODEL_NAME_FINALIZACAO
    elif GENERATIVE_MODEL_INSTANCE_GERACAO:
        modelo_a_usar = GENERATIVE_MODEL_INSTANCE_GERACAO
        nome_modelo_em_uso = GEMINI_MODEL_NAME_GERACAO
    
    if not GEMINI_API_AVAILABLE or modelo_a_usar is None:
        exibir_status(f"[SIMULAÇÃO - API INDISPONÍVEL/MODELO NÃO CONFIGURADO] [AGENTE {tipo_agente.upper()}] Processando solicitação...")
        return simular_resposta_fallback(prompt_para_ia, tipo_agente)

    exibir_status(f"[AGENTE {tipo_agente.upper()}] Enviando solicitação para o modelo '{nome_modelo_em_uso}'...", delay=1)
    
    try:
        response = modelo_a_usar.generate_content(prompt_para_ia)
        
        if hasattr(response, 'text'):
            return response.text
        elif response.parts:
            return "".join(part.text for part in response.parts if hasattr(part, 'text'))
        elif response.prompt_feedback and response.prompt_feedback.block_reason:
            block_reason = response.prompt_feedback.block_reason
            print(f"❌ A solicitação foi bloqueada pela API. Razão: {block_reason}")
            if response.prompt_feedback.safety_ratings:
                 for rating in response.prompt_feedback.safety_ratings:
                    print(f"   - Categoria: {rating.category}, Probabilidade: {rating.probability.name if hasattr(rating.probability, 'name') else rating.probability}")
            return f"ERRO_API: Solicitação bloqueada ({block_reason})."
        else:
            try:
                if response.candidates and response.candidates[0].content and response.candidates[0].content.parts:
                    return "".join(part.text for part in response.candidates[0].content.parts if hasattr(part, 'text'))
            except (AttributeError, IndexError): pass
            print(f"❌ A API Gemini retornou uma resposta inesperada. Resposta: {response}")
            return "ERRO_API: Resposta vazia ou inesperada."
    except Exception as e:
        print(f"❌ Erro durante a chamada à API Gemini: {e}")
        return f"ERRO_API: {e}."

def simular_resposta_fallback(prompt_para_ia, tipo_agente):
    """Fornece respostas simuladas se a API real não puder ser usada."""
    time.sleep(1)
    if "pesquisar os principais tópicos" in prompt_para_ia.lower() or tipo_agente == "pesquisa_esboco":
        return ("Relatório de Tópicos Sugeridos (Simulado):\n1. Introdução\n2. Desenvolvimento\n3. Conclusão")
    elif "pesquisar conteúdo para o capítulo" in prompt_para_ia.lower() or tipo_agente == "pesquisa_capitulo":
        return ("Conteúdo pesquisado simulado para o capítulo:\n- Ponto chave A\n- Ponto chave B")
    elif "escrever o conteúdo do capítulo" in prompt_para_ia.lower() or tipo_agente == "escritor":
        return ("Este é um parágrafo simulado gerado como fallback. A API Gemini não está configurada.")
    elif "revise o seguinte texto do capítulo" in prompt_para_ia.lower() or tipo_agente == "finalizador" or "TEXTO COMPLETO PARA REVISÃO FINAL" in prompt_para_ia:
        texto_original_match = re.search(r"TEXTO COMPLETO PARA REVISÃO FINAL:\n(.*?)$", prompt_para_ia, re.DOTALL)
        if texto_original_match:
            texto_original = texto_original_match.group(1).strip()
            return f"{texto_original}\n\n(Nota: Esta é uma revisão simulada. O texto foi retornado como estava.)"
        return "Texto (simuladamente) revisado e finalizado. (Nota: Revisão simulada sem extração de original)."
    else:
        return "Resposta simulada genérica (API não disponível)."


# --- Fase 1: Coleta de Informações e Planejamento ---

def coletar_tema_livro():
    exibir_cabecalho("Definição do Tema do Livro")
    return obter_input_usuario("Sobre qual tema você gostaria de criar um livro?")

def definir_esboco_inicial(tema_livro):
    exibir_cabecalho("Esboço Inicial do Livro")
    esboco_atual = ""
    while True:
        prompt_esboco = f"Estou planejando um livro sobre '{tema_livro}'. Pesquise e liste os principais tópicos, subtemas e áreas importantes que poderiam ser abordados em um livro sobre este tema. Apresente como um relatório conciso em formato de lista numerada ou com marcadores."
        sugestao_esboco = chamar_api_gemini(prompt_esboco, tipo_agente="pesquisa_esboco")

        print("\nSugestão de esboço gerada pela IA:")
        print("-" * 30); print(sugestao_esboco); print("-" * 30)

        if "ERRO_API" in sugestao_esboco:
            print("⚠️ Não foi possível gerar o esboço via IA. Por favor, insira manualmente.")
            linhas_esboco = []
            print("Digite o esboço do seu livro (tópicos principais, um por linha). Digite 'FIM_ESBOCO' em uma nova linha para terminar.")
            while True:
                linha = input()
                if linha.strip().upper() == 'FIM_ESBOCO': break
                linhas_esboco.append(linha)
            esboco_atual = "\n".join(linhas_esboco)
            if not esboco_atual.strip():
                print("❌ Esboço manual não pode ser vazio. Tente novamente."); continue
            return esboco_atual

        if not esboco_atual: esboco_atual = sugestao_esboco
        confirmacao = obter_confirmacao("Você aprova este esboço? (s/n - para nova sugestão / m - para modificar / a - para adicionar)", ('s','n','m','a'))
        if confirmacao == 's': print("✅ Esboço aprovado!"); return esboco_atual
        elif confirmacao == 'n': esboco_atual = ""; exibir_status("Ok, vamos gerar uma nova sugestão.")
        elif confirmacao == 'm':
            print("\nEdite o esboço. Digite 'FIM_EDICAO' em uma nova linha para terminar.")
            print("-" * 30); print(esboco_atual); print("-" * 30)
            linhas_editadas = []
            while True:
                linha = input()
                if linha.strip().upper() == 'FIM_EDICAO': break
                linhas_editadas.append(linha)
            esboco_atual = "\n".join(linhas_editadas)
            print("✅ Esboço modificado.")
            reconfirm = obter_confirmacao(f"Este é o esboço final?\n---\n{esboco_atual}\n---\n(s/n - para editar mais)", ('s','n'))
            if reconfirm == 's': print("✅ Esboço final aprovado!"); return esboco_atual
        elif confirmacao == 'a':
            adicao = obter_input_usuario("O que você gostaria de adicionar ao esboço atual?")
            esboco_atual += f"\n- {adicao} (Adicionado pelo usuário)"
            print("✅ Item adicionado. Esboço atualizado:"); print("-" * 30); print(esboco_atual); print("-" * 30)

def gerar_resumo_detalhado(esboco_livro, tema_livro):
    exibir_cabecalho("Resumo Detalhado do Livro")
    prompt_resumo = f"Com base no seguinte esboço aprovado para um livro sobre '{tema_livro}': \n{esboco_livro}\nCrie um resumo detalhado e coeso do que será abordado no livro. Este resumo servirá como contexto principal para a criação dos capítulos."
    resumo = chamar_api_gemini(prompt_resumo, tipo_agente="escritor")
    print("\nResumo Detalhado Gerado:"); print("-" * 30); print(resumo); print("-" * 30)

    if "ERRO_API" in resumo:
        print("⚠️ Não foi possível gerar o resumo via IA.")
        if obter_confirmacao("Deseja inserir um resumo manualmente? (s/n)", ('s','n')) == 's':
            linhas_resumo = []
            print("Digite o resumo. Digite 'FIM_RESUMO' em uma nova linha para terminar.")
            while True:
                linha = input()
                if linha.strip().upper() == 'FIM_RESUMO': break
                linhas_resumo.append(linha)
            resumo = "\n".join(linhas_resumo)
            if not resumo.strip(): return "Resumo manual não fornecido."
        else: return "Resumo não gerado devido a erro da API."

    if obter_confirmacao("Você aprova este resumo detalhado?", ('s','n')) == 's':
        print("✅ Resumo detalhado aprovado!"); return resumo
    else:
        print("ℹ️  Resumo não aprovado. Usaremos este como base."); return resumo

def definir_titulo_livro(contexto_geral_livro, tema_livro):
    exibir_cabecalho("Título do Livro")
    if obter_confirmacao("Gostaria que a IA gerasse sugestões de título?", ('s','n')) == 's':
        prompt_titulos = f"Sugira 5 títulos criativos para um livro com o tema '{tema_livro}' e resumo: \n'{contexto_geral_livro}'. Liste apenas os títulos, um por linha."
        sugestoes_raw = chamar_api_gemini(prompt_titulos, tipo_agente="geral")
        
        if "ERRO_API" in sugestoes_raw: print("⚠️ Não foi possível gerar sugestões de título via IA.")
        else:
            print("\nSugestões de Título:"); print("-" * 30); print(sugestoes_raw); print("-" * 30)
            titulos_lista = [re.sub(r"^\d+\.\s*", "", t.strip()) for t in sugestoes_raw.split('\n') if t.strip()]

            while True:
                if not titulos_lista: print("ℹ️ Nenhuma sugestão válida da IA."); break
                escolha_str = obter_input_usuario(f"Digite o número do título (1-{len(titulos_lista)} ou '0' para inserir o seu):")
                try:
                    escolha = int(escolha_str)
                    if 0 <= escolha <= len(titulos_lista):
                        if escolha == 0: titulo_final = obter_input_usuario("Digite o título do seu livro:")
                        else: titulo_final = titulos_lista[escolha - 1]
                        print(f"✅ Título definido: '{titulo_final}'"); return titulo_final
                    else: print(f"❌ Escolha inválida.")
                except ValueError: print("❌ Por favor, digite um número.")
    
    titulo_final = obter_input_usuario("Digite o título do seu livro:")
    print(f"✅ Título definido: '{titulo_final}'"); return titulo_final

def definir_quantidade_capitulos():
    exibir_cabecalho("Quantidade de Capítulos")
    return obter_input_usuario("Quantos capítulos o livro terá?", tipo_esperado=int, validacao_func=lambda x: x > 0, erro_msg="Deve ser maior que zero.")

def definir_titulos_capitulos(contexto_geral_livro, qtd_capitulos, titulo_livro):
    exibir_cabecalho("Títulos dos Capítulos")
    while True:
        prompt_tit_caps = (f"Para o livro '{titulo_livro}' ({qtd_capitulos} capítulos, resumo: '{contexto_geral_livro}'), "
                           f"sugira um título conciso para cada um dos {qtd_capitulos} capítulos. Liste um por linha, sem numeração.")
        sugestoes_raw = chamar_api_gemini(prompt_tit_caps, tipo_agente="geral")
        
        if "ERRO_API" in sugestoes_raw:
            print("⚠️ Não foi possível gerar títulos de capítulo via IA. Insira manualmente.")
            sugestoes_lista = [obter_input_usuario(f"Título para Capítulo {i+1}:") for i in range(qtd_capitulos)]
            return sugestoes_lista

        sugestoes_lista = [s.strip() for s in sugestoes_raw.split('\n') if s.strip() and not s.lower().startswith("aqui estão")]
        
        if len(sugestoes_lista) != qtd_capitulos:
            print(f"⚠️ IA sugeriu {len(sugestoes_lista)} títulos, pedimos {qtd_capitulos}. Ajustando/Completando...")
            if len(sugestoes_lista) > qtd_capitulos: sugestoes_lista = sugestoes_lista[:qtd_capitulos]
            else:
                for i in range(len(sugestoes_lista), qtd_capitulos):
                    sugestoes_lista.append(obter_input_usuario(f"Título para Capítulo {i+1} (faltante):"))
        
        print("\nSugestões de Títulos:"); print("-" * 30)
        for i, titulo_cap in enumerate(sugestoes_lista): print(f"{i+1}. {titulo_cap}"); print("-" * 30)

        confirmacao = obter_confirmacao("Aprova esta lista de títulos? (s/n - nova sugestão / m - modificar)", ('s','n','m'))
        if confirmacao == 's': print("✅ Títulos aprovados!"); return sugestoes_lista
        elif confirmacao == 'n': exibir_status("Ok, gerando nova lista...")
        elif confirmacao == 'm':
            titulos_finais_mod = list(sugestoes_lista)
            for i in range(len(titulos_finais_mod)):
                novo_titulo_cap = obter_input_usuario(f"Capítulo {i+1} (Enter para manter '{titulos_finais_mod[i]}'):")
                if novo_titulo_cap: titulos_finais_mod[i] = novo_titulo_cap
            print("✅ Títulos modificados."); return titulos_finais_mod

def definir_publico_alvo(contexto_geral_livro, titulo_livro):
    exibir_cabecalho("Público-Alvo do Livro")
    prompt_publico = (f"Para o livro '{titulo_livro}' (contexto: '{contexto_geral_livro}'), "
                      "liste exemplos de público-alvo em marcadores.")
    sugestoes = chamar_api_gemini(prompt_publico, tipo_agente="geral")
    
    if "ERRO_API" not in sugestoes: print("\nExemplos de Público-Alvo da IA:"); print("-" * 30); print(sugestoes); print("-" * 30)
    else: print("⚠️ Não foi possível obter sugestões de público-alvo da IA.")
    return obter_input_usuario("Qual o público-alvo principal do seu livro? (Descreva brevemente)")

def definir_limites_paragrafos():
    exibir_cabecalho("Limites de Parágrafos por Capítulo")
    min_p = obter_input_usuario("MÍNIMO de parágrafos por capítulo?", tipo_esperado=int, validacao_func=lambda x: x > 0, erro_msg="Mínimo > 0.")
    while True:
        max_p = obter_input_usuario("MÁXIMO de parágrafos por capítulo?", tipo_esperado=int, validacao_func=lambda x: x >= min_p, erro_msg=f"Máximo >= {min_p}.")
        if max_p >= min_p: return min_p, max_p
        else: print(f"❌ Máximo ({max_p}) não pode ser menor que mínimo ({min_p}).")

# --- Fase 2: Criação Iterativa dos Capítulos ---
def criar_capitulos(planejamento, modo_autonomo_capitulos=False):
    exibir_cabecalho("Criação dos Capítulos")
    if modo_autonomo_capitulos:
        print("🤖 MODO AUTÔNOMO ATIVADO PARA CRIAÇÃO DE CAPÍTULOS 🤖")
        exibir_status("A IA tomará decisões de aprovação e revisão automaticamente.", delay=1.5)

    livro_final_data = {
        "titulo_livro": planejamento['titulo_livro'], "tema_livro": planejamento['tema_livro'],
        "contexto_geral_livro": planejamento['contexto_geral_livro'],
        "publico_alvo": planejamento['publico_alvo'], "capitulos": []
    }
    resumos_caps_anteriores_para_contexto = []

    for i, titulo_cap_atual in enumerate(planejamento['titulos_capitulos']):
        if not modo_autonomo_capitulos: limpar_tela()
        exibir_cabecalho(f"Capítulo {i+1}/{len(planejamento['titulos_capitulos'])}: {titulo_cap_atual}")

        contexto_caps_anteriores_str = "\n".join(resumos_caps_anteriores_para_contexto)
        if not contexto_caps_anteriores_str: contexto_caps_anteriores_str = "Este é o primeiro capítulo."
        else: contexto_caps_anteriores_str = f"Contexto dos capítulos anteriores (resumos):\n{contexto_caps_anteriores_str}"

        conteudo_pesquisado_aprovado = ""
        while True: # Loop para pesquisa de conteúdo do capítulo
            prompt_pesquisa_cap = (
                f"Para o capítulo '{titulo_cap_atual}' do livro '{planejamento['titulo_livro']}' "
                f"(público-alvo: {planejamento['publico_alvo']}), cujo tema geral é '{planejamento['contexto_geral_livro']}', "
                f"e considerando o {contexto_caps_anteriores_str}. "
                "Pesquise e forneça os principais pontos, informações relevantes, dados e conceitos a serem abordados neste capítulo específico. Seja detalhado e forneça material substancial."
            )
            conteudo_pesquisado = chamar_api_gemini(prompt_pesquisa_cap, tipo_agente="pesquisa_capitulo")
            print("\nConteúdo/Tópicos para o capítulo (sugerido pela IA):"); print("-" * 30); print(conteudo_pesquisado); print("-" * 30)

            if "ERRO_API" in conteudo_pesquisado:
                print("⚠️ Não foi possível pesquisar conteúdo via IA.")
                if modo_autonomo_capitulos:
                    print("🤖 Modo Autônomo: Pulando pesquisa devido a erro."); conteudo_pesquisado_aprovado = "ERRO PESQUISA (AUTONOMO)"; break
                else:
                    print("Você precisará fornecer os pontos principais manualmente.")
                    linhas_conteudo = []; print("Digite os pontos/conteúdo. Digite 'FIM_CONTEUDO' para terminar.")
                    while True:
                        linha = input();
                        if linha.strip().upper() == 'FIM_CONTEUDO': break
                        linhas_conteudo.append(linha)
                    conteudo_pesquisado_aprovado = "\n".join(linhas_conteudo)
                    if not conteudo_pesquisado_aprovado.strip(): print("❌ Conteúdo manual não pode ser vazio."); continue
                    break
            
            if modo_autonomo_capitulos:
                print("🤖 Modo Autônomo: Aprovando conteúdo pesquisado."); conteudo_pesquisado_aprovado = conteudo_pesquisado; break
            else:
                confirmacao_pesquisa = obter_confirmacao("Aprova este conteúdo/tópicos? (s/n - nova pesquisa / r - refazer com sugestão)", ('s','n','r'))
                if confirmacao_pesquisa == 's': conteudo_pesquisado_aprovado = conteudo_pesquisado; print("✅ Conteúdo aprovado."); break
                elif confirmacao_pesquisa == 'n': exibir_status("Ok, nova pesquisa...")
                elif confirmacao_pesquisa == 'r':
                    sugestao_usuario = obter_input_usuario("Qual sua sugestão para refazer a pesquisa?")
                    exibir_status(f"Refazendo pesquisa com sugestão: '{sugestao_usuario}'...")
        
        texto_capitulo_final = ""
        while True: # Loop para geração e aprovação do texto do capítulo
            prompt_escrita = (
                f"Aja como um escritor especialista e professor experiente no tema do livro '{planejamento['titulo_livro']}'. "
                f"Escreva o conteúdo completo do capítulo '{titulo_cap_atual}'. "
                f"O público-alvo é: {planejamento['publico_alvo']}. "
                f"Baseie-se no seguinte material/tópicos: \n'{conteudo_pesquisado_aprovado}'.\n"
                f"O capítulo deve ter entre {planejamento['min_paragrafos']} e {planejamento['max_paragrafos']} parágrafos. "
                f"Mantenha um tom adequado. Considere o {contexto_caps_anteriores_str} para fluidez.\n"
                "Foque em clareza e coesão. Não adicione 'Neste capítulo...' a menos que natural."
            )
            texto_capitulo_gerado = chamar_api_gemini(prompt_escrita, tipo_agente="escritor")
            print("\nTexto do capítulo gerado pela IA:"); print("-" * 30); print(texto_capitulo_gerado); print("-" * 30)

            if "ERRO_API" in texto_capitulo_gerado:
                print("⚠️ Não foi possível gerar o texto do capítulo via IA.")
                if modo_autonomo_capitulos:
                    print("🤖 Modo Autônomo: Erro ao gerar texto."); texto_capitulo_final = "ERRO GERAÇÃO TEXTO (AUTONOMO)"; break
                elif obter_confirmacao("Tentar gerar novamente? (s/n)", ('s','n')) == 'n':
                    texto_capitulo_final = "ERRO GERAÇÃO CAPÍTULO."; break 
                else: continue

            revisar_pela_ia = modo_autonomo_capitulos or (obter_confirmacao("IA deve revisar este texto? (s/n)", ('s','n')) == 's')
            if modo_autonomo_capitulos and revisar_pela_ia: print("🤖 Modo Autônomo: Solicitando revisão automática.")

            if revisar_pela_ia:
                prompt_revisao = (
                    f"Revise o texto do capítulo '{titulo_cap_atual}' (livro: '{planejamento['titulo_livro']}', público: {planejamento['publico_alvo']}, "
                    f"contexto geral e caps anteriores: {planejamento['contexto_geral_livro']} {contexto_caps_anteriores_str}) "
                    f"quanto à gramática (Português), coesão, clareza. Texto para revisão: \n'{texto_capitulo_gerado}'\n"
                    "Forneça a versão revisada. Se poucas alterações, confirme."
                )
                texto_revisado_sugerido = chamar_api_gemini(prompt_revisao, tipo_agente="revisor")
                print("\nTexto revisado pela IA:"); print("-" * 30); print(texto_revisado_sugerido); print("-" * 30)
                
                if "ERRO_API" not in texto_revisado_sugerido:
                    if modo_autonomo_capitulos:
                        print("🤖 Modo Autônomo: Aplicando texto revisado."); texto_capitulo_gerado = texto_revisado_sugerido 
                    elif obter_confirmacao("Usar texto revisado pela IA? (s/n)", ('s','n')) == 's':
                        texto_capitulo_gerado = texto_revisado_sugerido; print("✅ Texto revisado aplicado.")
                    else: print("ℹ️  Texto original gerado mantido.")
                else: print("⚠️ Erro na revisão da IA. Texto original mantido.")

            if modo_autonomo_capitulos:
                print("🤖 Modo Autônomo: Aprovando texto do capítulo."); texto_capitulo_final = texto_capitulo_gerado; break
            else:
                confirmacao_escrita = obter_confirmacao("Aprova o texto do capítulo? (s/n - reescrever / m - pedir melhorias)", ('s','n','m'))
                if confirmacao_escrita == 's': texto_capitulo_final = texto_capitulo_gerado; print("✅ Texto aprovado!"); break
                elif confirmacao_escrita == 'n': exibir_status("Ok, gerando texto novamente...")
                elif confirmacao_escrita == 'm':
                    sugestao_melhora = obter_input_usuario("O que melhorar/alterar?")
                    exibir_status(f"Regerando com sugestão: '{sugestao_melhora}'...")
        
        resumo_para_contexto_agente_finalizador = "Resumo não disponível." # Default
        prompt_resumo_cap = (
            f"Resuma os pontos principais do capítulo '{titulo_cap_atual}' em 2-3 frases. "
            f"Contexto do capítulo:\n{texto_capitulo_final}"
        )
        resumo_cap_atual = chamar_api_gemini(prompt_resumo_cap, tipo_agente="geral")
        if "ERRO_API" not in resumo_cap_atual and resumo_cap_atual:
            resumo_extraido = resumo_cap_atual.split("Resumo do capítulo", 1)[-1].split(":", 1)[-1].strip() if ":" in resumo_cap_atual else resumo_cap_atual.strip()
            if not resumo_extraido: resumo_extraido = "Resumo não pôde ser extraído."
            resumos_caps_anteriores_para_contexto.append(f"Cap. {i+1} ({titulo_cap_atual}): {resumo_extraido}")
            resumo_para_contexto_agente_finalizador = resumo_extraido # Salvar para o agente finalizador
            print(f"💬 Resumo para contexto: {resumo_extraido}")
        else:
            print("⚠️ Não foi possível gerar resumo do capítulo."); resumos_caps_anteriores_para_contexto.append(f"Cap. {i+1} ({titulo_cap_atual}): Resumo não disponível.")

        capitulo_data = {
            "numero": i + 1, "titulo": titulo_cap_atual,
            "conteudo_pesquisado_aprovado": conteudo_pesquisado_aprovado,
            "texto_final": texto_capitulo_final,
            "resumo_para_contexto": resumo_para_contexto_agente_finalizador # ADICIONADO AQUI
        }
        livro_final_data["capitulos"].append(capitulo_data)
        salvar_progresso_livro_json(livro_final_data)
        exibir_status(f"Capítulo '{titulo_cap_atual}' salvo em '{NOME_ARQUIVO_PROGRESO}'.")


        if i < len(planejamento['titulos_capitulos']) - 1:
            if not modo_autonomo_capitulos:
                if obter_confirmacao("Prosseguir para o próximo capítulo?", ('s','n')) != 's':
                    print("⚠️ Criação interrompida."); break
            else: exibir_status(f"🤖 Modo Autônomo: Próximo capítulo ({i+2})...", delay=1.5)
        else: exibir_status("🎉 Todos os capítulos foram criados!")
    return livro_final_data

def salvar_progresso_livro_json(livro_data, nome_arquivo=NOME_ARQUIVO_PROGRESO):
    try:
        with open(nome_arquivo, 'w', encoding='utf-8') as f:
            json.dump(livro_data, f, ensure_ascii=False, indent=4)
    except IOError as e: print(f"❌ Erro ao salvar progresso: {e}")

# --- Agente Finalizador e Geração de Arquivos ---

def agente_finalizador(livro_data_completa, planejamento_completo):
    """
    Agente responsável pela revisão final de todos os capítulos e geração dos documentos.
    Utiliza o modelo de FINALIZAÇÃO configurado em uma ÚNICA chamada otimizada.
    """
    exibir_cabecalho("Agente Finalizador - Revisão Geral Otimizada")
    if not GENERATIVE_MODEL_INSTANCE_FINALIZACAO:
        print("⚠️ Modelo de finalização não configurado. Não é possível executar o Agente Finalizador.")
        print("   Gerando arquivos com o conteúdo atual...")
        if livro_data_completa and livro_data_completa.get("capitulos"):
            nome_base = sanitizar_nome_arquivo(livro_data_completa['titulo_livro'])
            if DOCX_AVAILABLE: gerar_arquivo_docx(livro_data_completa, nome_base)
            if REPORTLAB_AVAILABLE: gerar_arquivo_pdf(livro_data_completa, nome_base)
        return

    print(f"ℹ️  O Agente Finalizador usará o modelo '{GEMINI_MODEL_NAME_FINALIZACAO}' para a revisão holística do livro.")
    
    # 1. Montar o texto completo do livro para uma única chamada de API
    texto_completo_livro = ""
    for cap_info in livro_data_completa.get('capitulos', []):
        texto_original_cap = cap_info.get('texto_final', '')
        # Adicionar um separador claro entre os capítulos
        texto_completo_livro += f"--- CAPÍTULO {cap_info['numero']}: {cap_info['titulo']} ---\n"
        texto_completo_livro += f"{texto_original_cap}\n\n"

    if not texto_completo_livro.strip():
        print("❌ Não há conteúdo para ser revisado pelo Agente Finalizador.")
        return

    # 2. Montar o prompt único e otimizado
    contexto_prompt_final = (
        f"Contexto Geral do Livro (Tema: {planejamento_completo['tema_livro']}, "
        f"Título do Livro: {planejamento_completo['titulo_livro']}, "
        f"Público-Alvo: {planejamento_completo['publico_alvo']}):\n{planejamento_completo['contexto_geral_livro']}"
    )

    prompt_finalizacao_livro_inteiro = (
        f"Você é um editor de livros sênior e especialista em '{planejamento_completo['tema_livro']}'. "
        f"Realize uma revisão final e polimento no MANUSCRITO COMPLETO a seguir. "
        f"{contexto_prompt_final}\n\n"
        f"Instruções para Revisão Final Holística:\n"
        f"1. Gramática e Ortografia: Corrija todos os erros em Português.\n"
        f"2. Coesão e Coerência Global: Garanta que o livro inteiro flua bem, que as ideias estejam conectadas logicamente entre os capítulos e que não haja contradições.\n"
        f"3. Estilo e Tom Consistentes: Ajuste o texto para manter um tom consistente com o público-alvo ('{planejamento_completo['publico_alvo']}') em todo o livro.\n"
        f"4. Clareza, Precisão e Redundância: Melhore a clareza, corrija imprecisões e elimine repetições desnecessárias ao longo de todo o manuscrito.\n"
        f"5. Não altere fundamentalmente o significado ou os principais pontos dos capítulos, apenas refine-os para a mais alta qualidade editorial.\n"
        f"6. **IMPORTANTE**: Retorne o TEXTO COMPLETO E REVISADO do livro, mantendo EXATAMENTE a mesma estrutura de separadores de capítulo que você recebeu (--- CAPÍTULO X: [TÍTULO] ---). Sua resposta deve começar diretamente com '--- CAPÍTULO 1: ...' e terminar após o último capítulo, sem adicionar comentários ou introduções.\n\n"
        f"TEXTO COMPLETO PARA REVISÃO FINAL:\n{texto_completo_livro}"
    )

    exibir_status(f"Enviando manuscrito completo para revisão final com o modelo '{GEMINI_MODEL_NAME_FINALIZACAO}'...")
    livro_inteiro_revisado = chamar_api_gemini(prompt_finalizacao_livro_inteiro, tipo_agente="finalizador", usar_modelo_finalizacao=True)

    if "ERRO_API" in livro_inteiro_revisado or not livro_inteiro_revisado.strip():
        print(f"⚠️ Erro ao finalizar o livro com o Agente Finalizador. O conteúdo original será mantido.")
        livro_revisado_data = livro_data_completa # Mantém os dados originais se a revisão falhar
    else:
        print("✅ Livro revisado e finalizado pelo Agente.")
        # 3. Analisar (parse) a resposta e atualizar os dados do livro
        # Usando regex para dividir o texto pelos separadores de capítulo
        # O lookahead `(?=--- CAPÍTULO \d+:)` garante que o separador seja mantido no início de cada parte
        capitulos_revisados_texto = re.split(r'(?=--- CAPÍTULO \d+:)', livro_inteiro_revisado)
        capitulos_revisados_texto = [cap.strip() for cap in capitulos_revisados_texto if cap.strip()]

        if len(capitulos_revisados_texto) == len(livro_data_completa['capitulos']):
            for i, cap_revisado_completo in enumerate(capitulos_revisados_texto):
                # Remove a linha do título do texto do capítulo, já que o título está nos dados
                texto_apenas = re.sub(r'--- CAPÍTULO \d+:.*?---\n', '', cap_revisado_completo, count=1).strip()
                livro_data_completa['capitulos'][i]['texto_final'] = texto_apenas
            
            livro_revisado_data = livro_data_completa
            salvar_progresso_livro_json(livro_revisado_data, "livro_totalmente_revisado.json")
            exibir_status("💾 Dados do livro atualizados com a revisão final.")
        else:
            print(f"⚠️ Erro ao analisar a resposta do Agente Finalizador. O número de capítulos retornado ({len(capitulos_revisados_texto)}) é diferente do esperado ({len(livro_data_completa['capitulos'])}).")
            print("   O conteúdo original será mantido para evitar perda de dados.")
            livro_revisado_data = livro_data_completa
    
    # Geração dos arquivos com o livro completamente revisado (ou original, se a revisão falhou)
    if obter_confirmacao("\nDeseja gerar os arquivos DOCX e PDF com a versão final do livro?", ('s','n')) == 's':
        nome_base = sanitizar_nome_arquivo(livro_revisado_data['titulo_livro'])
        if DOCX_AVAILABLE: gerar_arquivo_docx(livro_revisado_data, nome_base)
        if REPORTLAB_AVAILABLE: gerar_arquivo_pdf(livro_revisado_data, nome_base)
    else:
        print("ℹ️  Geração dos arquivos finais pulada. O progresso finalizado está salvo nos arquivos JSON.")


def carregar_dados_livro_json(nome_arquivo=NOME_ARQUIVO_PROGRESO):
    try:
        with open(nome_arquivo, 'r', encoding='utf-8') as f: return json.load(f)
    except FileNotFoundError: print(f"❌ Arquivo '{nome_arquivo}' não encontrado."); return None
    except json.JSONDecodeError: print(f"❌ Erro ao decodificar JSON '{nome_arquivo}'."); return None

def sanitizar_nome_arquivo(nome):
    nome = re.sub(r'[^\w\s-]', '', nome).strip(); nome = re.sub(r'[-\s]+', '_', nome)
    return nome if nome else "livro_sem_titulo"

def gerar_arquivo_docx(livro_data, nome_arquivo_base_sanitizado):
    if not DOCX_AVAILABLE: print("ℹ️  Geração de DOCX pulada (biblioteca não disponível)."); return
    exibir_cabecalho("Geração do Arquivo DOCX")
    exibir_status(f"Gerando arquivo DOCX: {nome_arquivo_base_sanitizado}.docx ...")
    doc = Document()
    doc.styles['Title'].font.name = 'Arial'; doc.styles['Title'].font.size = Pt(28)
    doc.styles['Heading 1'].font.name = 'Arial'; doc.styles['Heading 1'].font.size = Pt(18)
    normal_style = doc.styles['Normal']; normal_style.font.name = 'Calibri'; normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5; normal_style.paragraph_format.space_after = Pt(12)

    titulo_livro_p = doc.add_paragraph(livro_data['titulo_livro'], style='Title')
    titulo_livro_p.alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_page_break()
    doc.add_heading('Sumário', level=1)
    for cap_info in livro_data.get('capitulos', []):
        doc.add_paragraph(f"Capítulo {cap_info['numero']}: {cap_info['titulo']}", style='ListBullet')
    doc.add_page_break()

    for cap_info in livro_data.get('capitulos', []):
        doc.add_heading(f"Capítulo {cap_info['numero']}: {cap_info['titulo']}", level=1)
        texto_final_cap = cap_info.get('texto_final', "Conteúdo não disponível.") 
        paragrafos_texto = texto_final_cap.split('\n\n') if texto_final_cap else ["Conteúdo não disponível."]
        for p_texto in paragrafos_texto:
            if p_texto.strip():
                paragrafo_doc = doc.add_paragraph(p_texto.strip(), style='Normal')
                paragrafo_doc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()
    nome_arquivo_docx = f"{nome_arquivo_base_sanitizado}.docx"
    try: doc.save(nome_arquivo_docx); print(f"✅ Livro salvo como '{nome_arquivo_docx}'")
    except Exception as e: print(f"❌ Erro ao salvar o arquivo DOCX: {e}")

def gerar_arquivo_pdf(livro_data, nome_arquivo_base_sanitizado):
    if not REPORTLAB_AVAILABLE: print("ℹ️  Geração de PDF pulada (biblioteca não disponível)."); return
    exibir_cabecalho("Geração do Arquivo PDF")
    exibir_status(f"Gerando arquivo PDF: {nome_arquivo_base_sanitizado}.pdf ...")
    nome_arquivo_pdf = f"{nome_arquivo_base_sanitizado}.pdf"
    doc_pdf = SimpleDocTemplate(nome_arquivo_pdf, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='TituloLivro', parent=styles['h1'], fontSize=28, alignment=TA_CENTER, spaceAfter=0.5*inch))
    styles.add(ParagraphStyle(name='TituloCapitulo', parent=styles['h2'], fontSize=18, spaceAfter=0.2*inch, spaceBefore=0.3*inch))
    styles.add(ParagraphStyle(name='CorpoTexto', parent=styles['Normal'], fontSize=12, alignment=TA_JUSTIFY, leading=14, spaceAfter=0.15*inch))
    styles.add(ParagraphStyle(name='SumarioItem', parent=styles['Normal'], fontSize=12, leftIndent=20, spaceAfter=0.1*inch))
    story = []
    story.append(Paragraph(livro_data['titulo_livro'], styles['TituloLivro'])); story.append(PageBreak())
    story.append(Paragraph("Sumário", styles['h1'])); story.append(Spacer(1, 0.2*inch))
    for cap_info in livro_data.get('capitulos', []):
        story.append(Paragraph(f"Capítulo {cap_info['numero']}: {cap_info['titulo']}", styles['SumarioItem']))
    story.append(PageBreak())
    for cap_info in livro_data.get('capitulos', []):
        story.append(Paragraph(f"Capítulo {cap_info['numero']}: {cap_info['titulo']}", styles['TituloCapitulo']))
        texto_final_cap = cap_info.get('texto_final', "Conteúdo não disponível.")
        paragrafos_texto = texto_final_cap.split('\n\n') if texto_final_cap else ["Conteúdo não disponível."]
        for p_texto in paragrafos_texto:
            if p_texto.strip(): story.append(Paragraph(p_texto.strip(), styles['CorpoTexto']))
        story.append(PageBreak())
    try: doc_pdf.build(story); print(f"✅ Livro salvo como '{nome_arquivo_pdf}'")
    except Exception as e: print(f"❌ Erro ao salvar o arquivo PDF: {e}")

# --- Função Principal (Orquestrador) ---
def main():
    limpar_tela()
    print("=" * 60); print("📚 Bem-vindo ao Assistente de Criação de Livros com IA! 📚".center(60)); print("=" * 60)
    
    api_configurada_com_sucesso = False
    if not GEMINI_API_AVAILABLE:
        print("\n‼️  A biblioteca 'google-generativeai' é essencial para este programa.")
        print("‼️  Por favor, instale-a com 'pip install google-generativeai' e tente novamente.")
        print("‼️  O programa continuará em modo de simulação limitada.")
    elif configurar_api_gemini():
        api_configurada_com_sucesso = True
    else: 
        print("\n‼️  Falha na configuração da API Gemini.")
        if obter_confirmacao("Deseja continuar em modo de simulação limitada? (s/n)", ('s','n')) != 's':
            print("👋 Saindo do programa.")
            return
        else:
            print("ℹ️  Continuando em modo de simulação...")
    
    print("\nVamos começar a planejar seu livro passo a passo.")

    tema_livro = coletar_tema_livro()
    esboco_livro = definir_esboco_inicial(tema_livro)
    contexto_geral_livro = gerar_resumo_detalhado(esboco_livro, tema_livro)
    if "ERRO_API" in contexto_geral_livro and "Resumo não gerado" in contexto_geral_livro : 
        print("❌ Criação do livro cancelada pois o resumo detalhado não pôde ser gerado ou inserido.")
        return

    titulo_livro = definir_titulo_livro(contexto_geral_livro, tema_livro)
    qtd_capitulos = definir_quantidade_capitulos()
    titulos_capitulos = definir_titulos_capitulos(contexto_geral_livro, qtd_capitulos, titulo_livro)
    publico_alvo = definir_publico_alvo(contexto_geral_livro, titulo_livro)
    min_paragrafos, max_paragrafos = definir_limites_paragrafos()

    planejamento_completo = {
        "tema_livro": tema_livro, "esboco_livro": esboco_livro, 
        "contexto_geral_livro": contexto_geral_livro, "titulo_livro": titulo_livro,
        "qtd_capitulos": qtd_capitulos, "titulos_capitulos": titulos_capitulos,
        "publico_alvo": publico_alvo, "min_paragrafos": min_paragrafos, "max_paragrafos": max_paragrafos
    }
    exibir_status("💾 Todas as informações de planejamento foram coletadas.")

    modo_autonomo_capitulos = False
    if api_configurada_com_sucesso:
        if obter_confirmacao("\nDeseja ativar o MODO AUTÔNOMO para a criação dos capítulos? (A IA aprovará e revisará automaticamente)", ('s','n')) == 's':
            modo_autonomo_capitulos = True
    
    livro_final_data = None # Inicializa
    if obter_confirmacao("\nPronto para iniciar a criação dos capítulos?", ('s','n')) == 's':
        livro_final_data = criar_capitulos(planejamento_completo, modo_autonomo_capitulos)
        
        if livro_final_data and livro_final_data.get("capitulos"):
            capitulos_com_conteudo_real = any(
                "ERRO NA GERAÇÃO" not in cap.get("texto_final", "").upper() and \
                "ERRO GERAÇÃO TEXTO" not in cap.get("texto_final", "").upper() and \
                "ERRO PESQUISA" not in cap.get("conteudo_pesquisado_aprovado", "").upper()
                for cap in livro_final_data["capitulos"]
            )
            if capitulos_com_conteudo_real:
                if api_configurada_com_sucesso and GENERATIVE_MODEL_INSTANCE_FINALIZACAO:
                    if obter_confirmacao("\nDeseja executar o AGENTE FINALIZADOR para revisão geral e geração dos documentos? (Recomendado)", ('s','n')) == 's':
                        agente_finalizador(livro_final_data, planejamento_completo)
                    else:
                        print("ℹ️  Agente Finalizador não executado. Gerando arquivos com o conteúdo atual (pós-criação de capítulos)...")
                        if obter_confirmacao("\nDeseja gerar os arquivos DOCX e PDF com o conteúdo atual?", ('s','n')) == 's':
                            nome_base = sanitizar_nome_arquivo(livro_final_data['titulo_livro'])
                            if DOCX_AVAILABLE: gerar_arquivo_docx(livro_final_data, nome_base)
                            if REPORTLAB_AVAILABLE: gerar_arquivo_pdf(livro_final_data, nome_base)
                else: 
                    print("ℹ️  API de finalização não configurada ou Agente Finalizador não executado.")
                    print("   Gerando arquivos com o conteúdo atual (pós-criação de capítulos)...")
                    if obter_confirmacao("\nDeseja gerar os arquivos DOCX e PDF com o conteúdo atual?", ('s','n')) == 's':
                        nome_base = sanitizar_nome_arquivo(livro_final_data['titulo_livro'])
                        if DOCX_AVAILABLE: gerar_arquivo_docx(livro_final_data, nome_base)
                        if REPORTLAB_AVAILABLE: gerar_arquivo_pdf(livro_final_data, nome_base)

            else:
                print(" ") 
                print("‼️  Geração de arquivos não realizada.")
                print("‼️  Motivo: Nenhum capítulo foi gerado com sucesso ou todos os capítulos contêm marcadores de erro.")
                print("‼️  Verifique o arquivo 'livro_em_progresso.json' para detalhes dos capítulos.")
                print("‼️  Você pode precisar revisar os capítulos manualmente ou tentar a geração novamente com prompts diferentes se houve falhas da API.")
        else: 
            print("ℹ️  Nenhum capítulo criado/aprovado. Nada para finalizar ou gerar.")
    else: 
        print("ℹ️  Criação dos capítulos cancelada.")

    print("\n" + "=" * 60); print("👋 Processo finalizado!".center(60)); print("=" * 60)

if __name__ == "__main__":
    if not GEMINI_API_AVAILABLE and (not DOCX_AVAILABLE or not REPORTLAB_AVAILABLE):
         print("*"*60)
         print("ATENÇÃO: Bibliotecas importantes não encontradas.")
         if not GEMINI_API_AVAILABLE: print("- google-generativeai (essencial para IA)")
         if not DOCX_AVAILABLE: print("- python-docx (para arquivos .docx)")
         if not REPORTLAB_AVAILABLE: print("- reportlab (para arquivos .pdf)")
         print("Por favor, instale-as para funcionalidade completa.")
         print("*"*60)
         input("Pressione Enter para tentar continuar...")
    main()
