# -*- coding: utf-8 -*-

from ui_utils import exibir_cabecalho, exibir_status
from file_utils import sanitizar_nome_arquivo

# --- Bibliotecas para Geração de Documentos ---
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

def aplicar_estilos_docx(doc, modelo):
    # Estilos base (podem ser sobrescritos ou estendidos por modelos)
    doc.styles['Title'].font.name = 'Arial'
    doc.styles['Title'].font.size = Pt(28)
    doc.styles['Heading 1'].font.name = 'Arial'
    doc.styles['Heading 1'].font.size = Pt(18)
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(12)

    if modelo == "abnt":
        doc.styles['Title'].font.name = 'Times New Roman'
        doc.styles['Title'].font.size = Pt(16)
        doc.styles['Title'].font.bold = True
        doc.styles['Heading 1'].font.name = 'Times New Roman'
        doc.styles['Heading 1'].font.size = Pt(14)
        doc.styles['Heading 1'].font.bold = True
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.line_spacing = 1.5
        normal_style.paragraph_format.first_line_indent = Inches(0.5)
        normal_style.paragraph_format.space_after = Pt(0)
        # Outras configurações ABNT podem ser adicionadas aqui (margens, etc. no objeto Document)
    elif modelo == "romance_moderno":
        doc.styles['Title'].font.name = 'Georgia'
        doc.styles['Title'].font.size = Pt(36)
        doc.styles['Heading 1'].font.name = 'Georgia'
        doc.styles['Heading 1'].font.size = Pt(22)
        doc.styles['Heading 1'].font.italic = True
        normal_style.font.name = 'Georgia'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.line_spacing = 1.8
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.first_line_indent = Inches(0.3)
    # Adicionar outros modelos aqui
    # else: (modelo padrão já aplicado no início)

def gerar_arquivo_docx(livro_data, nome_arquivo_base_sanitizado, modelo="padrao"):
    if not DOCX_AVAILABLE: print("ℹ️  Geração de DOCX pulada (biblioteca não disponível)."); return
    exibir_cabecalho("Geração do Arquivo DOCX")
    exibir_status(f"Gerando arquivo DOCX: {nome_arquivo_base_sanitizado}.docx ...")
    doc = Document()
    aplicar_estilos_docx(doc, modelo)

    # Configurações de página ABNT (exemplo, idealmente isso seria mais robusto)
    if modelo == "abnt":
        section = doc.sections[0]
        section.top_margin = Inches(1.18) # 3cm
        section.bottom_margin = Inches(0.78) # 2cm
        section.left_margin = Inches(1.18) # 3cm
        section.right_margin = Inches(0.78) # 2cm

    titulo_livro_p = doc.add_paragraph(livro_data['titulo_livro'], style='Title')
    titulo_livro_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Adicionar seções adicionais antes do sumário e capítulos
    secoes_adicionais = livro_data.get('secoes_adicionais', {})
    # Ordenar seções para uma ordem comum (ex: Prefácio, Introdução)
    ordem_preferencial = ['prefacio', 'introducao'] # Chaves sanitizadas
    secoes_ordenadas = sorted(secoes_adicionais.items(), key=lambda item: ordem_preferencial.index(item[0]) if item[0] in ordem_preferencial else len(ordem_preferencial))

    for nome_secao_sanitizado, secao_data in secoes_ordenadas:
        if secao_data.get('conteudo'):
            doc.add_heading(secao_data['titulo'], level=1)
            paragrafos_secao = secao_data['conteudo'].split('\n\n')
            for p_texto in paragrafos_secao:
                if p_texto.strip():
                    paragrafo_doc = doc.add_paragraph(p_texto.strip(), style='Normal')
                    paragrafo_doc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_page_break()
    doc.add_heading('Sumário', level=1)
    for cap_info in livro_data.get('capitulos', []):
        doc.add_paragraph(f"Capítulo {cap_info['numero']}: {cap_info['titulo']}", style='ListBullet')
    doc.add_page_break()

    for cap_info in livro_data.get('capitulos', []):
        doc.add_heading(f"Capítulo {cap_info['numero']}: {cap_info['titulo']}", level=1)
        texto_final_cap = cap_info.get('texto_final', "Conteúdo não disponível.") 
        paragrafos_texto = texto_final_cap.split('\n\n') if texto_final_cap else ["Conteúdo não disponível."]
        for p_texto in paragrafos_texto:
            if p_texto.strip():
                paragrafo_doc = doc.add_paragraph(p_texto.strip(), style='Normal')
                paragrafo_doc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()
    nome_arquivo_docx = f"{nome_arquivo_base_sanitizado}.docx"
    try: doc.save(nome_arquivo_docx); print(f"✅ Livro salvo como '{nome_arquivo_docx}'")
    except Exception as e: print(f"❌ Erro ao salvar o arquivo DOCX: {e}")

def obter_estilos_pdf(modelo):
    styles = getSampleStyleSheet()
    # Estilos base
    styles.add(ParagraphStyle(name='TituloLivro', parent=styles['h1'], fontSize=28, alignment=TA_CENTER, spaceAfter=0.5*inch))
    styles.add(ParagraphStyle(name='TituloCapitulo', parent=styles['h2'], fontSize=18, spaceAfter=0.2*inch, spaceBefore=0.3*inch))
    styles.add(ParagraphStyle(name='CorpoTexto', parent=styles['Normal'], fontSize=12, alignment=TA_JUSTIFY, leading=14, spaceAfter=0.15*inch))
    styles.add(ParagraphStyle(name='SumarioItem', parent=styles['Normal'], fontSize=12, leftIndent=20, spaceAfter=0.1*inch))

    if modelo == "abnt":
        styles['TituloLivro'].fontName = 'Times-Roman'
        styles['TituloLivro'].fontSize = 16
        styles['TituloLivro'].leading = 18
        styles['TituloLivro'].spaceAfter = 0.2*inch
        styles['TituloCapitulo'].fontName = 'Times-Roman'
        styles['TituloCapitulo'].fontSize = 14
        styles['TituloCapitulo'].leading = 16
        styles['CorpoTexto'].fontName = 'Times-Roman'
        styles['CorpoTexto'].fontSize = 12
        styles['CorpoTexto'].leading = 18 # Espaçamento 1.5
        styles['CorpoTexto'].firstLineIndent = 0.5*inch
        styles['CorpoTexto'].spaceAfter = 0
        styles['SumarioItem'].fontName = 'Times-Roman'
    elif modelo == "romance_moderno":
        styles['TituloLivro'].fontName = 'Georgia'
        styles['TituloLivro'].fontSize = 36
        styles['TituloCapitulo'].fontName = 'Georgia'
        styles['TituloCapitulo'].fontSize = 22
        # styles['TituloCapitulo'].fontStyle = 'Italic' # ReportLab não tem 'fontStyle' direto, precisa de <font name=Georgia-Italic>...</font>
        styles['CorpoTexto'].fontName = 'Georgia'
        styles['CorpoTexto'].fontSize = 11
        styles['CorpoTexto'].leading = 20 # Espaçamento maior
        styles['CorpoTexto'].firstLineIndent = 0.3*inch
    # Adicionar outros modelos aqui
    return styles

def gerar_arquivo_pdf(livro_data, nome_arquivo_base_sanitizado, modelo="padrao"):
    if not REPORTLAB_AVAILABLE: print("ℹ️  Geração de PDF pulada (biblioteca não disponível)."); return
    exibir_cabecalho("Geração do Arquivo PDF")
    exibir_status(f"Gerando arquivo PDF: {nome_arquivo_base_sanitizado}.pdf ...")
    nome_arquivo_pdf = f"{nome_arquivo_base_sanitizado}.pdf"
    styles = obter_estilos_pdf(modelo)
    # Configurações de página ABNT (exemplo)
    if modelo == "abnt":
        doc_pdf = SimpleDocTemplate(nome_arquivo_pdf, pagesize=letter, 
                                    rightMargin=0.78*inch, leftMargin=1.18*inch, 
                                    topMargin=1.18*inch, bottomMargin=0.78*inch)
    else:
        doc_pdf = SimpleDocTemplate(nome_arquivo_pdf, pagesize=letter, 
                                    rightMargin=inch, leftMargin=inch, 
                                    topMargin=inch, bottomMargin=inch)
    story = []
    story.append(Paragraph(livro_data['titulo_livro'], styles['TituloLivro']))
    story.append(PageBreak())

    # Adicionar seções adicionais antes do sumário e capítulos
    secoes_adicionais = livro_data.get('secoes_adicionais', {})
    ordem_preferencial = ['prefacio', 'introducao'] # Chaves sanitizadas
    secoes_ordenadas = sorted(secoes_adicionais.items(), key=lambda item: ordem_preferencial.index(item[0]) if item[0] in ordem_preferencial else len(ordem_preferencial))

    for nome_secao_sanitizado, secao_data in secoes_ordenadas:
        if secao_data.get('conteudo'):
            story.append(Paragraph(secao_data['titulo'], styles['TituloCapitulo'])) # Usar estilo de capítulo para seções
            paragrafos_secao = secao_data['conteudo'].split('\n\n')
            for p_texto in paragrafos_secao:
                if p_texto.strip():
                    story.append(Paragraph(p_texto.strip(), styles['CorpoTexto']))
            story.append(PageBreak())
    story.append(Paragraph("Sumário", styles['h1'])); story.append(Spacer(1, 0.2*inch))
    for cap_info in livro_data.get('capitulos', []):
        story.append(Paragraph(f"Capítulo {cap_info['numero']}: {cap_info['titulo']}", styles['SumarioItem']))
    story.append(PageBreak())
    for cap_info in livro_data.get('capitulos', []):
        story.append(Paragraph(f"Capítulo {cap_info['numero']}: {cap_info['titulo']}", styles['TituloCapitulo']))
        texto_final_cap = cap_info.get('texto_final', "Conteúdo não disponível.")
        paragrafos_texto = texto_final_cap.split('\n\n') if texto_final_cap else ["Conteúdo não disponível."]
        for p_texto in paragrafos_texto:
            if p_texto.strip(): story.append(Paragraph(p_texto.strip(), styles['CorpoTexto']))
        story.append(PageBreak())
    try: doc_pdf.build(story); print(f"✅ Livro salvo como '{nome_arquivo_pdf}'")
    except Exception as e: print(f"❌ Erro ao salvar o arquivo PDF: {e}")